import os
import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_border(cell, **kwargs):
    """
    Set cell border
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    
    for key, value in kwargs.items():
        if value:
            tag = 'w:{}'.format(key)
            element = OxmlElement(tag)
            element.set(qn('w:val'), value)
            tcPr.append(element)

def add_hyperlink(paragraph, text, url):
    """
    Add a hyperlink to a paragraph
    """
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id)
    
    new_run = docx.oxml.shared.OxmlElement('w:r')
    rPr = docx.oxml.shared.OxmlElement('w:rPr')
    
    # Add styling to the hyperlink
    c = docx.oxml.shared.OxmlElement('w:color')
    c.set(docx.oxml.shared.qn('w:val'), '0000FF')
    rPr.append(c)
    
    u = docx.oxml.shared.OxmlElement('w:u')
    u.set(docx.oxml.shared.qn('w:val'), 'single')
    rPr.append(u)
    
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    
    paragraph._p.append(hyperlink)
    return hyperlink

def format_cv_from_markdown(markdown_file, output_docx):
    """
    Convert markdown CV to professionally formatted DOCX
    """
    # Read markdown content
    with open(markdown_file, 'r') as f:
        content = f.read()
    
    # Split content by sections
    sections = content.split('## ')
    
    # Create a new document
    doc = docx.Document()
    
    # Set document margins
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
    
    # Define styles
    styles = doc.styles
    
    # Heading 1 style (for name)
    h1_style = styles.add_style('Heading1Custom', WD_STYLE_TYPE.PARAGRAPH)
    h1_style.base_style = styles['Heading 1']
    h1_font = h1_style.font
    h1_font.name = 'Calibri'
    h1_font.size = Pt(16)
    h1_font.bold = True
    h1_font.color.rgb = RGBColor(0, 0, 0)
    
    # Heading 2 style (for section headers)
    h2_style = styles.add_style('Heading2Custom', WD_STYLE_TYPE.PARAGRAPH)
    h2_style.base_style = styles['Heading 2']
    h2_font = h2_style.font
    h2_font.name = 'Calibri'
    h2_font.size = Pt(12)
    h2_font.bold = True
    h2_font.color.rgb = RGBColor(0, 51, 102)
    
    # Normal text style
    normal_style = styles['Normal']
    normal_font = normal_style.font
    normal_font.name = 'Calibri'
    normal_font.size = Pt(10)
    
    # Parse the markdown and create the document
    # First section is the title, which starts with # not ##
    title_section = sections[0].split('\n', 1)
    title = title_section[0].replace('# ', '')
    
    # Add name as title
    title_para = doc.add_paragraph(style='Heading1Custom')
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(title)
    
    # Process contact information
    contact_info = ''
    for line in title_section[1].strip().split('\n'):
        if line.startswith('## '):
            break
        if line.strip():
            contact_info += line.strip() + '\n'
    
    # Add contact information in a table
    contact_table = doc.add_table(rows=1, cols=3)
    contact_table.style = 'Table Grid'
    contact_table.autofit = True
    
    # Split contact info into three columns
    contact_lines = contact_info.strip().split('\n')
    cell_texts = []
    
    # Phone and Email in first cell
    cell_texts.append('\n'.join(contact_lines[:2]))
    
    # LinkedIn in second cell
    if len(contact_lines) > 2:
        cell_texts.append(contact_lines[2])
    else:
        cell_texts.append('')
    
    # Location and Citizenship in third cell
    if len(contact_lines) > 3:
        cell_texts.append('\n'.join(contact_lines[3:]))
    else:
        cell_texts.append('')
    
    # Fill the table cells
    for i, text in enumerate(cell_texts):
        cell = contact_table.cell(0, i)
        cell_para = cell.paragraphs[0]
        
        # Check if this is the LinkedIn cell
        if i == 1 and 'linkedin.com' in text:
            linkedin_parts = text.split('linkedin.com/')
            if len(linkedin_parts) > 1:
                linkedin_url = 'https://www.linkedin.com/' + linkedin_parts[1]
                add_hyperlink(cell_para, text, linkedin_url)
        else:
            cell_para.add_run(text)
        
        cell_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Set cell borders
        set_cell_border(cell, top="single", bottom="single", left="single", right="single")
    
    # Add a small space after contact info
    doc.add_paragraph()
    
    # Process the rest of the sections
    for section in sections[1:]:
        if not section.strip():
            continue
        
        # Split the section into title and content
        section_parts = section.split('\n', 1)
        section_title = section_parts[0].strip()
        
        # Add section header
        section_header = doc.add_paragraph(style='Heading2Custom')
        section_header.add_run(section_title.upper())
        
        # Add horizontal line below section header
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 0.1
        p.add_run('_' * 120)
        
        # Process section content
        if len(section_parts) > 1:
            section_content = section_parts[1].strip()
            
            # Special handling for different sections
            if section_title == 'PROFESSIONAL EXPERIENCE':
                # Split by company entries (### markers)
                company_entries = section_content.split('### ')
                
                for entry in company_entries[1:]:  # Skip the first empty split
                    entry_parts = entry.split('\n', 1)
                    company_title = entry_parts[0].strip()
                    
                    # Add company and position
                    company_para = doc.add_paragraph()
                    company_run = company_para.add_run(company_title)
                    company_run.bold = True
                    
                    # Add the rest of the entry content
                    if len(entry_parts) > 1:
                        entry_content = entry_parts[1].strip()
                        
                        # Process bullet points
                        for line in entry_content.split('\n'):
                            line = line.strip()
                            if line.startswith('- '):
                                # This is a bullet point
                                bullet_para = doc.add_paragraph(style='List Bullet')
                                bullet_para.add_run(line[2:])
                            elif line.startswith('**Achievement:**'):
                                # This is an achievement line
                                achievement_para = doc.add_paragraph()
                                achievement_text = line.replace('**Achievement:**', '').strip()
                                achievement_run = achievement_para.add_run('Achievement: ' + achievement_text)
                                achievement_run.italic = True
                                achievement_run.bold = True
                            elif line:
                                # Regular paragraph
                                regular_para = doc.add_paragraph()
                                regular_para.add_run(line)
                
            elif section_title == 'CORE COMPETENCIES':
                # Create a 2-column table for competencies
                competencies = [comp.strip()[2:] for comp in section_content.split('\n') if comp.strip().startswith('- ')]
                
                # Calculate rows needed (3 competencies per row)
                rows_needed = (len(competencies) + 2) // 3
                comp_table = doc.add_table(rows=rows_needed, cols=3)
                comp_table.style = 'Table Grid'
                comp_table.autofit = True
                
                # Fill the table with competencies
                for i, comp in enumerate(competencies):
                    row = i // 3
                    col = i % 3
                    cell = comp_table.cell(row, col)
                    cell_para = cell.paragraphs[0]
                    cell_para.add_run('• ' + comp)
                    
                    # Set cell borders (no borders for cleaner look)
                    set_cell_border(cell, top="nil", bottom="nil", left="nil", right="nil")
            
            elif section_title == 'EDUCATION' or section_title == 'TECHNICAL SKILLS' or section_title == 'PROFESSIONAL DEVELOPMENT':
                # Process education entries
                for line in section_content.split('\n'):
                    line = line.strip()
                    if line.startswith('### '):
                        # Institution name
                        institution = line.replace('### ', '')
                        inst_para = doc.add_paragraph()
                        inst_run = inst_para.add_run(institution)
                        inst_run.bold = True
                    elif line.startswith('**'):
                        # Degree information
                        degree_line = line.replace('**', '')
                        degree_para = doc.add_paragraph()
                        degree_para.paragraph_format.left_indent = Inches(0.2)
                        degree_para.add_run(degree_line)
                    elif line.startswith('- '):
                        # Bullet point
                        bullet_para = doc.add_paragraph(style='List Bullet')
                        bullet_para.add_run(line[2:])
                    elif line:
                        # Regular paragraph
                        regular_para = doc.add_paragraph()
                        regular_para.add_run(line)
            
            else:
                # Default handling for other sections
                p = doc.add_paragraph()
                p.add_run(section_content)
    
    # Save the document
    doc.save(output_docx)
    return output_docx

def convert_all_cvs():
    """
    Convert all markdown CVs to DOCX format
    """
    input_dir = '/home/ubuntu/cv_project/final_cvs'
    output_dir = '/home/ubuntu/cv_project/final_cvs/docx'
    
    # Create output directory if it doesn't exist
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
    
    # Process all markdown files
    for filename in os.listdir(input_dir):
        if filename.endswith('.md'):
            input_path = os.path.join(input_dir, filename)
            output_path = os.path.join(output_dir, filename.replace('.md', '.docx'))
            
            print(f"Converting {filename} to DOCX format...")
            format_cv_from_markdown(input_path, output_path)
            print(f"Created {output_path}")
    
    print("All CVs have been converted to DOCX format.")

if __name__ == "__main__":
    convert_all_cvs()
